from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify
from flask_cors import CORS
import fitz
import docx
import pandas as pd
import requests
import io
import os
import threading
import time
import uuid
from supabase import create_client, Client

app = Flask(__name__)
CORS(app)

OPENROUTER_KEY = os.environ.get("OPENROUTER_KEY", "")
SUPABASE_URL   = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY   = os.environ.get("SUPABASE_KEY", "")
COHERE_KEY     = os.environ.get("COHERE_KEY", "")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)


# ── COHERE EMBEDDING ─────────────────────────────────────────

def get_embedding(text, input_type="search_document"):
    try:
        response = requests.post(
            "https://api.cohere.com/v1/embed",
            headers={
                "Authorization": f"Bearer {COHERE_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "texts": [text[:2000]],  # ~500 tokens, safely within Cohere's 512 token limit
                "model": "embed-english-v3.0",
                "input_type": input_type
            }
        )
        result = response.json()
        if "embeddings" in result:
            return result["embeddings"][0]
        else:
            print(f"Cohere error: {result}")
            return None
    except Exception as e:
        print(f"Cohere embedding error: {e}")
        return None


# ── TEXT EXTRACTION ──────────────────────────────────────────
# extract_blocks() returns a list of blocks instead of one flat string.
# Each block is {"type": "text"|"table", "content": str}.
# This keeps tables intact end-to-end, even when they span multiple pages,
# so they never get split apart by the chunker later.

def _table_to_text(rows):
    """Convert a list of table rows (list of cells) into a clean text block."""
    lines = []
    for row in rows:
        cells = [str(c).replace("\n", " ").strip() if c is not None else "" for c in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def _tables_compatible(a_header, b_header):
    """Check if two tables likely belong to the same multi-page table (same column count)."""
    return len(a_header) == len(b_header) and len(a_header) > 0


def extract_blocks_from_pdf(file_bytes):
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    blocks = []
    pending_table = None  # holds {"header": [...], "rows": [...]} across pages

    for page in pdf:
        tabs = page.find_tables()
        table_bboxes = [fitz.Rect(t.bbox) for t in tabs.tables]

        # Get text only from regions NOT covered by a detected table,
        # so table content isn't duplicated as plain text.
        if table_bboxes:
            page_rect = page.rect
            text_clips = []
            # Above the first table
            text_clips.append(fitz.Rect(page_rect.x0, page_rect.y0, page_rect.x1, table_bboxes[0].y0))
            # Between tables
            for j in range(len(table_bboxes) - 1):
                text_clips.append(fitz.Rect(page_rect.x0, table_bboxes[j].y1, page_rect.x1, table_bboxes[j + 1].y0))
            # Below the last table
            text_clips.append(fitz.Rect(page_rect.x0, table_bboxes[-1].y1, page_rect.x1, page_rect.y1))

            page_text = ""
            for clip in text_clips:
                if clip.height > 2:  # skip degenerate/empty slivers
                    page_text += page.get_text(clip=clip)
        else:
            page_text = page.get_text()

        if tabs.tables:
            for t in tabs.tables:
                data = t.extract()
                if not data:
                    continue
                header, rows = data[0], data[1:]

                if pending_table and _tables_compatible(pending_table["header"], header):
                    # Looks like a continuation of the previous page's table
                    pending_table["rows"].extend(rows if rows else [header])
                else:
                    # Flush previous pending table as its own block
                    if pending_table:
                        full_rows = [pending_table["header"]] + pending_table["rows"]
                        blocks.append({"type": "table", "content": _table_to_text(full_rows)})
                    pending_table = {"header": header, "rows": rows}
        else:
            # No table on this page — flush any pending table first
            if pending_table:
                full_rows = [pending_table["header"]] + pending_table["rows"]
                blocks.append({"type": "table", "content": _table_to_text(full_rows)})
                pending_table = None

        if page_text.strip():
            blocks.append({"type": "text", "content": page_text.strip()})

    if pending_table:
        full_rows = [pending_table["header"]] + pending_table["rows"]
        blocks.append({"type": "table", "content": _table_to_text(full_rows)})

    pdf.close()
    return blocks


def extract_blocks_from_file(file):
    """Returns a list of {"type": "text"|"table", "content": str} blocks for any supported file."""
    filename = file.filename.lower()
    try:
        if filename.endswith(".txt") or filename.endswith(".md"):
            content = file.read().decode("utf-8", errors="ignore")
            return [{"type": "text", "content": content.strip()}] if content.strip() else []

        elif filename.endswith(".pdf"):
            file_bytes = file.read()
            return extract_blocks_from_pdf(file_bytes)

        elif filename.endswith(".docx"):
            file_bytes = file.read()
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n".join(para.text for para in doc.paragraphs)
            return [{"type": "text", "content": content.strip()}] if content.strip() else []

        elif filename.endswith(".csv"):
            file_bytes = file.read()
            df = pd.read_csv(io.BytesIO(file_bytes))
            content = df.to_string(index=False)
            return [{"type": "table", "content": content.strip()}] if content.strip() else []

        else:
            content = file.read().decode("utf-8", errors="ignore")
            return [{"type": "text", "content": content.strip()}] if content.strip() else []

    except Exception as e:
        return [{"type": "text", "content": f"Could not read file: {str(e)}"}]


# Kept for backward compatibility with any other code that expects flat text
def extract_text_from_file(file):
    blocks = extract_blocks_from_file(file)
    return "\n\n".join(b["content"] for b in blocks).strip()


# ── CHUNKING ─────────────────────────────────────────────────

def chunk_text(text, chunk_size=1000, overlap=100):
    """Word-count chunking with overlap — used only for 'text' type blocks."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def chunk_blocks(blocks, chunk_size=1000, overlap=100, max_table_chars=6000):
    """
    Turns extraction blocks into final chunks:
    - table blocks are kept whole (one chunk), even if they span pages,
      UNLESS they exceed max_table_chars — then split by row groups to stay safe.
    - text blocks go through normal word-count chunking with overlap.
    """
    chunks = []
    for block in blocks:
        if block["type"] == "table":
            content = block["content"]
            if len(content) <= max_table_chars:
                chunks.append(content)
            else:
                # Very large table — split by rows, not mid-row, to stay safe on token limits
                rows = content.split("\n")
                current, size = [], 0
                for row in rows:
                    if size + len(row) > max_table_chars and current:
                        chunks.append("\n".join(current))
                        current, size = [], 0
                    current.append(row)
                    size += len(row)
                if current:
                    chunks.append("\n".join(current))
        else:
            chunks.extend(chunk_text(block["content"], chunk_size, overlap))
    return chunks


# ── ROUTES ───────────────────────────────────────────────────

@app.route("/")
def home():
    return "Project QA Tool Backend is running!"


@app.route("/create_project", methods=["POST"])
def create_project():
    data = request.json
    name        = data.get("name", "")
    description = data.get("description", "")
    date        = data.get("date", "")

    if not name:
        return jsonify({"error": "Project name required"}), 400

    try:
        result = supabase.table("projects").insert({
            "name": name,
            "description": description,
            "date": date
        }).execute()
        return jsonify({"project": result.data[0]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/get_projects", methods=["GET"])
def get_projects():
    try:
        result = supabase.table("projects").select("*").order("created_at", desc=True).execute()
        return jsonify({"projects": result.data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/delete_project", methods=["POST"])
def delete_project():
    data = request.json
    project_id = data.get("project_id", "")
    try:
        supabase.table("projects").delete().eq("id", project_id).execute()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/upload", methods=["POST"])
def upload():
    files      = request.files.getlist("files")
    project_id = request.form.get("project_id", "")

    if not files:
        return jsonify({"error": "No files uploaded"}), 400
    if not project_id:
        return jsonify({"error": "Project ID required"}), 400

    uploaded = []

    for file in files:
        filename = file.filename
        blocks   = extract_blocks_from_file(file)
        if not blocks:
            continue

        chunks  = chunk_blocks(blocks, chunk_size=1000, overlap=100)
        records = []

        for chunk in chunks:
            embedding = get_embedding(chunk, input_type="search_document")
            if embedding:
                records.append({
                    "project_id": project_id,
                    "file_name":  filename,
                    "content":    chunk,
                    "embedding":  embedding
                })

        if records:
            supabase.table("chunks").insert(records).execute()
            uploaded.append({"name": filename, "chunks": len(records)})

    return jsonify({"uploaded": uploaded})


@app.route("/get_files", methods=["POST"])
def get_files():
    data       = request.json
    project_id = data.get("project_id", "")
    try:
        result = supabase.table("chunks").select("file_name").eq("project_id", project_id).execute()
        files  = list(set([r["file_name"] for r in result.data]))
        return jsonify({"files": files})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/delete_file", methods=["POST"])
def delete_file():
    data       = request.json
    project_id = data.get("project_id", "")
    file_name  = data.get("file_name", "")
    try:
        supabase.table("chunks").delete().eq("project_id", project_id).eq("file_name", file_name).execute()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/ask", methods=["POST"])
def ask():
    data         = request.json
    question     = data.get("question", "")
    project_id   = data.get("project_id", "")
    project_name = data.get("project_name", "")
    project_desc = data.get("project_desc", "")
    file_name    = data.get("file_name", None)
    history      = data.get("history", [])

    if not question:
        return jsonify({"error": "No question provided"}), 400

    query_embedding = get_embedding(question, input_type="search_query")
    if not query_embedding:
        return jsonify({"error": "Could not generate embedding for question"}), 500

    try:
        if file_name:
            result = supabase.rpc("match_chunks", {
                "query_embedding":  query_embedding,
                "match_project_id": project_id,
                "match_file_name":  file_name,
                "match_count":      5
            }).execute()
        else:
            result = supabase.rpc("match_chunks_all", {
                "query_embedding":  query_embedding,
                "match_project_id": project_id,
                "match_count":      5
            }).execute()

        matched_chunks = result.data
        print(f"Question: {question}")
        print(f"Matched chunks: {len(matched_chunks)}")
        for chunk in matched_chunks:
            print(f"  - {chunk['file_name']}: {chunk['content'][:100]}")

    except Exception as e:
        return jsonify({"error": f"Search error: {str(e)}"}), 500

    if not matched_chunks:
        return jsonify({"reply": "I couldn't find relevant information in the documents for your question.", "tokens_used": 0})

    context = ""
    for chunk in matched_chunks:
        context += f"\n--- From: {chunk['file_name']} ---\n"
        context += chunk['content']
        context += "\n"

    if file_name:
        system_msg = f"""You are an AI assistant for the project "{project_name}".
You are focused on the file: {file_name}

Here are the most relevant sections from this file:

{context}

Instructions:
1. Answer based ONLY on the provided sections.
2. Be precise and helpful.
3. If the answer is not in the sections, say so clearly."""
    else:
        system_msg = f"""You are an AI assistant for the project "{project_name}".
Project description: {project_desc}

Here are the most relevant sections from the project documents:

{context}

Instructions:
1. Answer based ONLY on the provided sections.
2. Always mention WHICH FILE the information comes from.
3. If asked to compare files, compare them clearly.
4. If the answer is not in the sections, say so clearly.
5. Format comparisons as bullet points or tables."""

    messages = [{"role": "system", "content": system_msg}]
    for msg in history:
        messages.append(msg)
    messages.append({"role": "user", "content": question})

    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_KEY}",
                "Content-Type":  "application/json",
                "HTTP-Referer":  "http://localhost:5000",
                "X-Title":       "Project QA Tool"
            },
            json={
                "model":      "openrouter/auto",
                "messages":   messages,
                "max_tokens": 2000
            },
            timeout=30
        )
        result = response.json()
        print(f"OpenRouter status: {response.status_code}")
        print(f"OpenRouter response: {result}")

        if "error" in result:
            error_msg = result["error"] if isinstance(result["error"], str) else result["error"].get("message", str(result["error"]))
            print(f"OpenRouter error: {error_msg}")
            return jsonify({"error": error_msg}), 500

        if "choices" not in result or not result["choices"]:
            print(f"No choices in response: {result}")
            return jsonify({"error": "No response from AI model. Try again."}), 500

        reply = result["choices"][0]["message"]["content"]

        # Get token usage from OpenRouter response
        tokens_used = 0
        if "usage" in result:
            tokens_used = result["usage"].get("total_tokens", 0)

        return jsonify({"reply": reply, "tokens_used": tokens_used})

    except Exception as e:
        print(f"OpenRouter exception: {e}")
        return jsonify({"error": str(e)}), 500


# ── CHAT HISTORY WITH SESSION SUPPORT ────────────────────────

@app.route("/save_message", methods=["POST", "OPTIONS"])
def save_message():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    data       = request.json
    project_id = data.get("project_id")
    role       = data.get("role")
    content    = data.get("content")
    session_id = data.get("session_id")
    tokens     = data.get("tokens_used", 0)
    try:
        supabase.table("chat_history").insert({
            "project_id": project_id,
            "role":       role,
            "content":    content,
            "session_id": session_id,
            "is_active":  True,
            "tokens_used": tokens
        }).execute()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/get_history", methods=["POST", "OPTIONS"])
def get_history():
    if request.method == "OPTIONS":
        return jsonify({}), 200
    data       = request.json
    project_id = data.get("project_id")
    session_id = data.get("session_id", None)
    try:
        query = supabase.table("chat_history").select("*")\
            .eq("project_id", project_id)\
            .eq("is_active", True)

        if session_id:
            query = query.eq("session_id", session_id)

        result = query.order("created_at").execute()
        return jsonify({"history": result.data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/new_chat", methods=["POST", "OPTIONS"])
def new_chat():
    """Soft delete - marks old messages as inactive, returns new session ID"""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    data       = request.json
    project_id = data.get("project_id")
    try:
        # Mark all current messages as inactive (soft delete)
        supabase.table("chat_history")\
            .update({"is_active": False})\
            .eq("project_id", project_id)\
            .eq("is_active", True)\
            .execute()

        # Generate new session ID
        new_session_id = str(uuid.uuid4())
        return jsonify({"success": True, "session_id": new_session_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/get_token_usage", methods=["POST", "OPTIONS"])
def get_token_usage():
    """Get token usage stats for a project"""
    if request.method == "OPTIONS":
        return jsonify({}), 200
    data       = request.json
    project_id = data.get("project_id")
    try:
        result = supabase.table("chat_history")\
            .select("role, content, tokens_used, created_at, session_id")\
            .eq("project_id", project_id)\
            .eq("role", "user")\
            .order("created_at", desc=True)\
            .execute()
        return jsonify({"usage": result.data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── KEEP ALIVE ───────────────────────────────────────────────

def keep_alive():
    while True:
        time.sleep(840)
        try:
            requests.get("https://project-qa-tool.onrender.com")
            print("Keep alive ping sent")
        except:
            pass


if __name__ == "__main__":
    t = threading.Thread(target=keep_alive)
    t.daemon = True
    t.start()
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)